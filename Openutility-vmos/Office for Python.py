try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
    from tkinter.simpledialog import askstring
    from PIL import Image, ImageTk
    from docx import Document
    from openpyxl import Workbook
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import datetime
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "tkinterweb"])
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pygments"])
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pillow"])

class SimpleOfficeApp:
    def __init__(self, root):
        
        # 注册中文字体（需准备字体文件）
        try:
            pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
        except:
            messagebox.showwarning("字体警告", "未找到中文字体文件，PDF导出可能显示异常")

        # 增强表格功能
        self.table_data = []
        self.current_cell = None
        self.create_table_toolbar()
        
        # 扩展文件菜单
        self.add_pdf_export_menu()

    def add_pdf_export_menu(self):
        """添加PDF导出菜单"""
        file_menu = self.root.nametowidget("!menu.!menu")
        file_menu.insert_command(4, label="导出为PDF", command=self.export_to_pdf)
        file_menu.insert(5, "separator")

    def create_table_toolbar(self):
        """创建表格专用工具栏"""
        table_toolbar = ttk.Frame(self.table_frame)
        
        table_buttons = [
            ("添加行", self.add_table_row),
            ("删除行", self.delete_table_row),
            ("添加列", self.add_table_column),
            ("删除列", self.delete_table_column),
            ("导出Excel", self.export_table_to_excel)
        ]
        
        for text, cmd in table_buttons:
            btn = ttk.Button(table_toolbar, text=text, command=cmd)
            btn.pack(side=tk.LEFT, padx=2)
        
        table_toolbar.grid(row=2, column=0, columnspan=2, sticky="ew")
        
        # 绑定单元格编辑事件
        self.table.bind("<Double-1>", self.on_cell_double_click)

    def on_cell_double_click(self, event):
        """处理单元格双击编辑"""
        region = self.table.identify_region(event.x, event.y)
        if region != "cell":
            return
        
        # 获取选中单元格信息
        column = self.table.identify_column(event.x)
        col_index = int(column[1:]) - 1
        item = self.table.identify_row(event.y)
        
        # 创建编辑框
        x, y, width, height = self.table.bbox(item, column)
        current_value = self.table.item(item, "values")[col_index]
        
        self.entry_edit = ttk.Entry(self.table_frame)
        self.entry_edit.place(x=x, y=y, width=width, height=height)
        self.entry_edit.insert(0, current_value)
        self.entry_edit.focus_set()
        
        # 绑定完成编辑事件
        self.entry_edit.bind("<FocusOut>", lambda e: self.save_cell_edit(item, col_index))
        self.entry_edit.bind("<Return>", lambda e: self.save_cell_edit(item, col_index))

    def save_cell_edit(self, item, col_index):
        """保存单元格编辑结果"""
        new_value = self.entry_edit.get()
        current_values = list(self.table.item(item, "values"))
        current_values[col_index] = new_value
        self.table.item(item, values=current_values)
        self.entry_edit.destroy()
        self.unsaved_changes = True

    def add_table_row(self):
        """添加新行"""
        columns = self.table["columns"]
        default_values = [f"新行{len(self.table.get_children())+1}"] + [""]*(len(columns)-1)
        self.table.insert("", tk.END, values=default_values)
        self.unsaved_changes = True

    def delete_table_row(self):
        """删除选中行"""
        selected = self.table.selection()
        if selected:
            self.table.delete(selected)
            self.unsaved_changes = True

    def add_table_column(self):
        """添加新列"""
        new_col = askstring("新建列", "输入列名称：")
        if new_col:
            columns = list(self.table["columns"])
            columns.append(new_col)
            self.table["columns"] = columns
            self.table.heading(new_col, text=new_col)
            self.table.column(new_col, width=100)
            self.unsaved_changes = True

    def delete_table_column(self):
        """删除选中列"""
        selected_col = self.table.focus()
        if selected_col:
            columns = list(self.table["columns"])
            if len(columns) > 1:
                columns.remove(selected_col)
                self.table["columns"] = columns
                self.unsaved_changes = True

    def export_table_to_excel(self):
        """导出表格到Excel"""
        if not self.table.get_children():
            messagebox.showwarning("警告", "表格为空，无法导出")
            return
        
        filepath = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel文件", "*.xlsx")]
        )
        if filepath:
            try:
                wb = Workbook()
                ws = wb.active
                
                # 写入表头
                headers = self.table["columns"]
                ws.append(headers)
                
                # 写入数据
                for child in self.table.get_children():
                    values = self.table.item(child, "values")
                    ws.append(values)
                
                wb.save(filepath)
                messagebox.showinfo("成功", "Excel文件导出成功")
            except Exception as e:
                messagebox.showerror("错误", f"导出失败：{str(e)}")

    def export_to_pdf(self):
        """导出文本内容到PDF"""
        text_content = self.text_editor.get("1.0", tk.END)
        if not text_content.strip():
            messagebox.showwarning("警告", "文档内容为空")
            return
        
        filepath = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF文件", "*.pdf")]
        )
        if filepath:
            try:
                c = canvas.Canvas(filepath, pagesize=A4)
                width, height = A4
                
                # 设置字体（需要中文字体支持）
                try:
                    c.setFont("SimSun", 12)
                except:
                    c.setFont("Helvetica", 12)
                
                # 设置边距
                margin = inch
                x = margin
                y = height - margin
                line_height = 14
                
                # 处理文本换行
                lines = []
                for line in text_content.split("\n"):
                    words = line.split()
                    current_line = []
                    for word in words:
                        if c.stringWidth(" ".join(current_line + [word])) < (width - 2*margin):
                            current_line.append(word)
                        else:
                            lines.append(" ".join(current_line))
                            current_line = [word]
                    if current_line:
                        lines.append(" ".join(current_line))
                
                # 写入PDF
                for line in lines:
                    if y < margin:
                        c.showPage()
                        y = height - margin
                        try:
                            c.setFont("SimSun", 12)
                        except:
                            c.setFont("Helvetica", 12)
                    c.drawString(x, y, line)
                    y -= line_height
                
                c.save()
                messagebox.showinfo("成功", "PDF导出成功")
            except Exception as e:
                messagebox.showerror("错误", f"PDF导出失败：{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = SimpleOfficeApp(root)
    root.mainloop()
